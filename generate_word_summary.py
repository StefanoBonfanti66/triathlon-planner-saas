from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Riassunto Preparazione Call CUS ProPatria', 0)

p = doc.add_paragraph()
p.add_run('Data Call: ').bold = True
p.add_run('Martedi 10 Marzo 2026\n')
p.add_run('Referente: ').bold = True
p.add_run('Alessandro Sciarrone')

doc.add_heading('1. Punti Chiave da Presentare', level=1)
doc.add_paragraph('Bulk Import (Excel): Caricamento massivo dei soci in pochi secondi.', style='List Bullet')
doc.add_paragraph('Log Attivita: Tracciabilita totale di ogni operazione admin per massima sicurezza.', style='List Bullet')
doc.add_paragraph('PWA & Branding: Icona e colori CUS ProPatria personalizzati per la squadra.', style='List Bullet')
doc.add_paragraph('Social Card Hub: Generazione automatica di card per Instagram.', style='List Bullet')

doc.add_heading('2. Checklist per la Demo Live', level=1)
doc.add_paragraph('Avviare ambiente locale (npm run dev) sul branch develop.', style='List Number')
doc.add_paragraph('Mostrare il pulsante "Importa Atleti" e caricare il file test_athletes_import.xlsx.', style='List Number')
doc.add_paragraph('Verificare il log dell\'operazione nella scheda "Log Attivita".', style='List Number')
doc.add_paragraph('Promuovere Alessandro ad Admin (cliccando sullo scudetto accanto al suo nome).', style='List Number')

doc.add_heading('3. Materiale Consegnato', level=1)
doc.add_paragraph('File Excel di Test: test_athletes_import.xlsx (nella root del progetto).')
doc.add_paragraph('Guida Admin PDF: GUIDA_ADMIN_CUS.pdf (copiata sul desktop).')

doc.save('C:/Users/localadm/Desktop/RIASSUNTO_CALL_CUS.docx')
print("✅ File Word creato con successo sul desktop!")
